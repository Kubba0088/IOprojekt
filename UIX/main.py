import os
import pyaudio
import json
import time
from vosk import Model, KaldiRecognizer
from kivy.app import App
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.properties import StringProperty
from kivy.lang import Builder
from threading import Thread
from kivy.clock import Clock
import pyautogui
import pytesseract
from PIL import Image
from fpdf import FPDF
from docx import Document
import pygetwindow as gw
from pywinauto import Desktop
from pywinauto import Application
from PIL import ImageGrab

model_path_pl = r"C:\Users\Barte\OneDrive\Pulpit\VOSK\vosk-model-small-pl-0.22" #zmien tutaj ścieżkę
model_path_en = r"C:\Users\Barte\OneDrive\Pulpit\VOSK\vosk-model-small-en-us-0.15" #zmien tutaj ścieżkę

model_pl = Model(model_path_pl)
model_en = Model(model_path_en)

output_file_path = r"C:\Users\Barte\OneDrive\Pulpit\Speach\speach.txt" #zmien tutaj ścieżkę
output_format = "txt"
current_model = model_pl
recognizer = KaldiRecognizer(current_model, 16000)
p = pyaudio.PyAudio()

input_device_index = 1

final_transcription = ""
last_speech_time = time.time()

PAUSE_THRESHOLD = 3

Builder.load_file("main_screen.kv")
Builder.load_file("transcription_screen.kv")
Builder.load_file("settings_screen.kv")

def get_window_bounds(app_name):
    try:
        app = Application(backend="win32").connect(title_re=app_name)
        window = app.window(title_re=app_name)
        rect = window.rectangle()  # Pobiera pozycję i wymiary okna
        return rect.left, rect.top, rect.right, rect.bottom
    except Exception as e:
        print(f"Błąd podczas pobierania wymiarów okna: {e}")
        return None
    
def capture_window_area(bounds):
    if bounds:
        left, top, right, bottom = bounds
        screenshot = ImageGrab.grab(bbox=(left, top, right, bottom))  # Zrzut obszaru
        return screenshot
    else:
        print("Nie udało się pobrać wymiarów okna.")
        return None

def debug_active_window():
    try:
        app = Application(backend="win32").connect(active_only=True)
        active_window = app.windows()[0]  # Pobierz pierwsze aktywne okno
        window_title = active_window.window_text()
        print(f"Debug: Tytuł okna: {window_title}")
        return window_title
    except Exception as e:
        print(f"Błąd podczas debugowania aktywnego okna: {e}")
        return None


def detect_meeting_app():
    try:
        app = Application(backend="win32").connect(active_only=True)  # Ogranicz do aktywnych aplikacji
        windows = app.windows()
        print("-------------------------------------------------")
        print(windows)
        for win in windows:
            window_title = win.window_text().lower()
            print(f"Debug: Tytuł okna - {window_title}")  # Debugowanie

            if "teams" in window_title:
                return "Microsoft Teams"
            elif "webex" in window_title:
                return "Webex"
        
        return "Nie wykryto aplikacji spotkań"
    except Exception as e:
        print(f"Błąd podczas wykrywania aplikacji: {e}")
        return "Nie można wykryć aplikacji"



    
def save_transcription_to_file(text):
    try:
        if output_format == "txt":
            with open(output_file_path, 'a', encoding='utf-8') as file:
                file.write(text + "\n")
            print(f"Transkrypcja zapisana jako TXT: {output_file_path}")
        elif output_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, text)
            pdf_file_path = output_file_path.replace(".txt", ".pdf")
            pdf.output(pdf_file_path)
            print(f"Transkrypcja zapisana jako PDF: {pdf_file_path}")
        elif output_format == "docx":
            doc = Document()
            doc.add_paragraph(text)
            docx_file_path = output_file_path.replace(".txt", ".docx")
            doc.save(docx_file_path)
            print(f"Transkrypcja zapisana jako DOCX: {docx_file_path}")
        else:
            print(f"Nieobsługiwany format zapisu: {output_format}")
    except Exception as e:
        print(f"Błąd podczas zapisywania do pliku: {e}")


class TeamsScreenMonitor:
    def __init__(self, output_dir=None, ocr_lang="eng", screenshot_prefix="teams_screenshot"):
        self.output_dir = output_dir if output_dir else os.path.join(os.path.expanduser("~"), "Desktop")
        self.ocr_lang = ocr_lang  # Dynamiczny język OCR
        self.screenshot_prefix = screenshot_prefix
        self.saved_texts = set()
        self.screen_number = 1

    def update_ocr_language(self, new_lang):
        """Aktualizuje język OCR."""
        print(f"Zmiana języka OCR na: {new_lang}")
        self.ocr_lang = new_lang

    def capture_fullscreen(self):
        try:
            screenshot = pyautogui.screenshot()
            screenshot_path = os.path.join(self.output_dir, f"{self.screenshot_prefix}_{self.screen_number}.png")
            screenshot.save(screenshot_path)
            print(f"Zrzut ekranu zapisany: {screenshot_path}")
            return screenshot_path
        except Exception as e:
            print(f"Błąd podczas robienia zrzutu ekranu: {e}")
            return None

    def perform_ocr_on_window(app_name):
        bounds = get_window_bounds(app_name)
        screenshot = capture_window_area(bounds)
        
        if screenshot:
            text = pytesseract.image_to_string(screenshot, lang='pol')  # OCR na zrzucie ekranu
            return text
        else:
            return "Nie udało się wykonać OCR."

    def detect_and_ocr_meeting_app():
        detected_app = detect_meeting_app()  # Funkcja, która wykrywa aplikację
        if detected_app in ["Microsoft Teams", "Webex"]:
            print(f"Przeprowadzanie OCR dla {detected_app}...")
            ocr_text = perform_ocr_on_window(detected_app)
            print(f"OCR wynik: {ocr_text}")
        else:
            print("Brak odpowiedniej aplikacji do wykonania OCR.")


    def perform_ocr(self, image_path):
        try:
            if not os.path.exists(image_path):
                print(f"Plik obrazu nie istnieje: {image_path}")
                return ""

            print(f"Otwieranie pliku obrazu: {image_path}")
            raw_text = pytesseract.image_to_string(Image.open(image_path), lang=self.ocr_lang)
            filtered_text = " ".join([word for word in raw_text.split() if len(word) > 2])
            return filtered_text
        except Exception as e:
            print(f"Błąd podczas wykonywania OCR: {e}")
            return ""

    def save_text_to_file(self, text):
        try:
            if output_format == "txt":
                text_file_path = os.path.join(self.output_dir, "recognized_text.txt")
                with open(text_file_path, 'a', encoding='utf-8') as file:
                    file.write(f"\n{'-' * 40}\nZrzut ekranu z {time.ctime()}:\n{text}")
                print(f"Rozpoznany tekst zapisany jako TXT: {text_file_path}")
            elif output_format == "pdf":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, text)
                pdf_file_path = os.path.join(self.output_dir, "recognized_text.pdf")
                pdf.output(pdf_file_path)
                print(f"Rozpoznany tekst zapisany jako PDF: {pdf_file_path}")
            elif output_format == "docx":
                doc = Document()
                doc.add_paragraph(text)
                docx_file_path = os.path.join(self.output_dir, "recognized_text.docx")
                doc.save(docx_file_path)
                print(f"Rozpoznany tekst zapisany jako DOCX: {docx_file_path}")
            else:
                print(f"Nieobsługiwany format zapisu: {output_format}")
        except Exception as e:
            print(f"Błąd podczas zapisywania tekstu z OCR: {e}")


    def process_screen(self):
        print("Robię zrzut i wykonuję OCR...")
        screenshot_path = self.capture_fullscreen()
        if screenshot_path:
            recognized_text = self.perform_ocr(screenshot_path)
            if recognized_text:
                if recognized_text not in self.saved_texts:
                    print(f"Rozpoznany tekst:\n{recognized_text}")
                    self.save_text_to_file(recognized_text)  # Wywołanie metody zapisu z dynamicznym formatem
                    self.saved_texts.add(recognized_text)
                else:
                    print("Ten tekst już został zapisany wcześniej, pomijam.")
            else:
                print("Nie udało się rozpoznać tekstu.")
        else:
            print("Nie udało się zrobić zrzutu ekranu.")

        self.screen_number += 1

    def start_monitoring(self, continue_ocr, sleep_interval=20):
        time.sleep(10)
        while continue_ocr():
            self.process_screen()
            time.sleep(sleep_interval)


class MainScreen(Screen):
    ocr_running = False
    current_app = StringProperty("Wykrywanie aplikacji...")  # Domyślny tekst

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.monitor_running = True
        self.monitor_instance = TeamsScreenMonitor(
            output_dir=r"C:\Users\Barte\OneDrive\Pulpit\Screenshots",
            ocr_lang="eng"
        )
        self.app_monitor_thread = Thread(target=self.monitor_meeting_application)
        self.app_monitor_thread.daemon = True
        self.app_monitor_thread.start()

    def monitor_meeting_application(self):
        while self.monitor_running:
            app_name = detect_meeting_app()
            print(f"Debug: Aktualnie wykryta aplikacja - {app_name}")
            self.current_app = f"Aktualna aplikacja: {app_name}"  # Aktualizuje tekst wyświetlany w GUI
            time.sleep(5)

    def on_leave(self):
        self.monitor_running = False

    def detect_meeting_application(self):
        app_name = detect_meeting_app()
        self.current_app = f"Aktualna aplikacja: {app_name}"
        print(self.current_app)

    def start_transcription(self):
        transcription_screen = self.manager.get_screen('transcription')
        transcription_screen.transcription_text = "Transkrypcja audio w toku..."

        transcription_thread = Thread(target=self.transcribe_audio, args=(transcription_screen,))
        transcription_thread.daemon = True
        transcription_thread.start()

    def start_ocr(self):
        transcription_screen = self.manager.get_screen('transcription')
        transcription_screen.transcription_text += "\nOCR w toku..."
        self.ocr_running = True

        monitor = TeamsScreenMonitor(output_dir=r"C:\Users\Barte\OneDrive\Pulpit\Screenshots",ocr_lang="eng") 
        monitor_thread = Thread(target=monitor.start_monitoring, args=(lambda: self.ocr_running,))
        monitor_thread.daemon = True
        monitor_thread.start()

    def stop_ocr(self):
        self.ocr_running = False

    def transcribe_audio(self, transcription_screen):
        global final_transcription, last_speech_time

        try:
            stream = p.open(format=pyaudio.paInt16, channels=1, rate=16000,
                            input=True, input_device_index=input_device_index,
                            frames_per_buffer=8000)
        except Exception as e:
            error_message = f"Nie można otworzyć strumienia audio: {e}"
            print(error_message)
            Clock.schedule_once(lambda dt: setattr(transcription_screen, 'transcription_text', error_message))
            return

        print("Nasłuchuję...")

        while True:
            try:
                data = stream.read(4000, exception_on_overflow=False)
            except Exception as e:
                error_message = f"Błąd podczas czytania strumienia audio: {e}"
                print(error_message)
                Clock.schedule_once(lambda dt: setattr(transcription_screen, 'transcription_text', error_message))
                continue

            if len(data) == 0:
                continue

            if recognizer.AcceptWaveform(data):
                result = recognizer.Result()
                result_json = json.loads(result)

                if 'text' in result_json and result_json['text']:
                    final_transcription += " " + result_json['text']
                    Clock.schedule_once(
                        lambda dt, text=final_transcription: self.update_transcription(transcription_screen, text))
                    last_speech_time = time.time()

            else:
                result = recognizer.PartialResult()
                result_json = json.loads(result)

                if 'partial' in result_json and result_json['partial']:
                    partial_text = result_json['partial']
                    Clock.schedule_once(
                        lambda dt, text=partial_text: self.update_partial_transcription(transcription_screen, text))
                    last_speech_time = time.time()

            if time.time() - last_speech_time > 5:
                if final_transcription.strip():
                    final_transcription += "."
                    Clock.schedule_once(
                        lambda dt, text=final_transcription.strip(): self.display_final_transcription(
                            transcription_screen, text))

                    save_transcription_to_file(final_transcription.strip())

                    final_transcription = ""
                last_speech_time = time.time()

    def update_transcription(self, transcription_screen, text):
        transcription_screen.transcription_text = text

    def update_partial_transcription(self, transcription_screen, text):
        transcription_screen.transcription_text = text

    def display_final_transcription(self, transcription_screen, text):
        transcription_screen.transcription_text = text


class TranscriptionScreen(Screen):
    transcription_text = StringProperty("Tu pojawi się podgląd transkrypcji.")


class SettingsScreen(Screen):
    def save_settings(self):
        global current_model, recognizer, output_format

        selected_language = self.ids.language_spinner.text
        selected_format = self.ids.format_spinner.text
        if selected_language == "polski":
            current_model = model_pl
            ocr_lang = "pol"  
        elif selected_language == "angielski":
            current_model = model_en
            ocr_lang = "eng"  

        recognizer = KaldiRecognizer(current_model, 16000)

   
        output_format = selected_format.lower()

        if hasattr(self.manager.get_screen('main'), 'monitor_instance'):
            monitor_instance = self.manager.get_screen('main').monitor_instance
            monitor_instance.update_ocr_language(ocr_lang)

        print(f"Zmieniono język na: {selected_language}")
        print(f"Zmieniono format zapisu na: {output_format}")

class MeetNotesApp(App):
    def build(self):
        sm = ScreenManager()
        sm.add_widget(MainScreen(name='main'))
        sm.add_widget(TranscriptionScreen(name='transcription'))
        sm.add_widget(SettingsScreen(name='settings'))
        return sm


if __name__ == "__main__":
    MeetNotesApp().run()
